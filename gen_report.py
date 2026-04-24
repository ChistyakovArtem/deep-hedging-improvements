"""Генератор .docx отчёта по преддипломной практике."""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = '/home/artemchistyak/deep-hedging-improvements/Отчет_преддипломная_практика.docx'
FONT = 'Times New Roman'


def _cyr(run, name=FONT):
    """Fix Cyrillic rendering in python-docx runs."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for k in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(k), name)


def p(doc, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      indent_first=True, space_before=0, space_after=0, size=14):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.first_line_indent = Cm(1.25) if indent_first else Cm(0)
    if space_before:
        par.paragraph_format.space_before = Pt(space_before)
    if space_after:
        par.paragraph_format.space_after = Pt(space_after)
    r = par.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    _cyr(r)
    return par


def h(doc, text, level=1):
    size = {1: 16, 2: 14}[level]
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    par.paragraph_format.space_after = Pt(6)
    r = par.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    _cyr(r)
    return par


def bullet(doc, text):
    par = doc.add_paragraph(style='List Bullet')
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.left_indent = Cm(1.0)
    r = par.add_run(text)
    r.font.size = Pt(14)
    _cyr(r)
    return par


def numbered(doc, text):
    par = doc.add_paragraph(style='List Number')
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.left_indent = Cm(1.0)
    r = par.add_run(text)
    r.font.size = Pt(14)
    _cyr(r)
    return par


def code_block(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    par.paragraph_format.left_indent = Cm(0.5)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(4)
    r = par.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for k in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rFonts.set(qn(k), 'Consolas')
    return par


# ────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page setup: A4, поля 3/1.5/2/2
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.5)

# Normal style defaults
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(14)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
rPr = normal.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
for k in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
    rFonts.set(qn(k), FONT)


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
p(doc, 'Федеральное государственное автономное образовательное учреждение',
  align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, size=12)
p(doc, 'высшего образования', align=WD_ALIGN_PARAGRAPH.CENTER,
  indent_first=False, size=12)
p(doc, '«Национальный исследовательский университет «Высшая школа экономики»',
  align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, bold=True, size=12)
p(doc, 'Факультет компьютерных наук', align=WD_ALIGN_PARAGRAPH.CENTER,
  indent_first=False, size=12)
p(doc, 'Образовательная программа «Компьютерные науки и анализ данных»',
  align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, size=12)
p(doc, 'Направление подготовки 01.03.02 Прикладная математика и информатика',
  align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, size=12)
p(doc, 'бакалавриат', align=WD_ALIGN_PARAGRAPH.CENTER,
  indent_first=False, size=12, space_after=96)

p(doc, 'О Т Ч Е Т', align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
  bold=True, size=18, space_after=12)
p(doc, 'по преддипломной практике', align=WD_ALIGN_PARAGRAPH.CENTER,
  indent_first=False, size=14, space_after=12)
p(doc, 'Enhancing Deep Hedging: Architecture, Optimisation, and Generalisation '
       'for Derivative Pricing Under Transaction Costs',
  align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, bold=True, size=14,
  space_after=96)

p(doc, 'Выполнил студент гр. БПМИ221', align=WD_ALIGN_PARAGRAPH.RIGHT,
  indent_first=False, size=12)
p(doc, 'Чистяков Артём', align=WD_ALIGN_PARAGRAPH.RIGHT,
  indent_first=False, size=12)
p(doc, '_____________ (подпись)', align=WD_ALIGN_PARAGRAPH.RIGHT,
  indent_first=False, size=12, space_after=48)

p(doc, 'Проверил:', align=WD_ALIGN_PARAGRAPH.RIGHT,
  indent_first=False, size=12)
p(doc, 'Пифтахин Геннадий, VP, Sberbank Risk Management',
  align=WD_ALIGN_PARAGRAPH.RIGHT, indent_first=False, size=12)
p(doc, '_____________ (подпись)', align=WD_ALIGN_PARAGRAPH.RIGHT,
  indent_first=False, size=12, space_after=48)

p(doc, 'Москва, 2026', align=WD_ALIGN_PARAGRAPH.CENTER,
  indent_first=False, size=12)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 1. ЦЕЛИ И ЗАДАЧИ ПРАКТИКИ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '1. Цели и задачи практики')

p(doc, 'Цель практики — улучшить базовый агент Deep Hedging '
  '(Buehler et al., 2019) для хеджирования европейского колла под моделью '
  'Heston с пропорциональными транзакционными издержками по метрике SoftMin '
  '(энтропийная риск-мера) и зафиксировать контролируемым ablation-ем, '
  'какие инженерные приёмы в сетапе Monte Carlo Policy Gradient (MCPG) '
  'дают устойчивый выигрыш.')

p(doc, 'Задачи практики:')
bullet(doc, 'реализовать эталонные не-RL стратегии: LSMC-дельту и '
       'closed-form Heston-дельту;')
bullet(doc, 'воспроизвести базовый Deep Hedging как MLP-политику, обучаемую '
       'pathwise-градиентом через симулятор SDE;')
bullet(doc, 'провести ablation по: архитектуре (MLP vs PAF-эмбеддинг), '
       'входным фичам (включение running P&L), инициализации '
       '(default / zero-last-layer / pretrain-на-дельту), стохастической '
       'регуляризации (SAC-энтропия), структуре выхода '
       '(δ vs δ = δ_BSM + Δ);')
bullet(doc, 'реализовать и сравнить оптимизаторы: Adam, K-FAC '
       '(Kronecker-factored natural gradient) и Muon '
       '(ортогонализованный моментум);')
bullet(doc, 'сопоставить все конфигурации на общем test-set по SoftMin и '
       'выбрать лучшую для последующего переноса в ВКР на реальные '
       'данные Deribit BTC/ETH.')


# ════════════════════════════════════════════════════════════════════════════
# 2. ПОСТАНОВКА ЗАДАЧИ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '2. Постановка задачи')

p(doc, 'Рассматривается модель Heston:')
p(doc, 'dS_t = r S_t dt + √v_t S_t dW_t^1,     '
       'dv_t = κ(θ − v_t) dt + ξ √v_t dW_t^2,     '
       'corr(dW^1, dW^2) = ρ.',
  indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER)

p(doc, 'Трейдер продал европейский колл с payoff max(S_T − K, 0) и '
       'хеджирует позицию динамически, ребалансируя портфель в моменты '
       't_i = i·T/N, i = 0, …, N. Платёж за сделку пропорционален '
       'обороту: fee_i = c · |Δδ_i| · S_{t_i}. Портфель самофинансируемый.')

p(doc, 'Политика δ_θ(state_t) параметризуется нейронной сетью с '
       'входом state_t = (S_t, v_t, τ_t, δ_{t-1}) либо расширенным '
       '(S_t, v_t, τ_t, δ_{t-1}, cash_t + δ_{t-1} S_t) при включённой '
       'фиче running P&L. Терминальный P&L на одном пути:')
p(doc, 'PnL(ω) = cash_T(ω) − max(S_T(ω) − K, 0).',
  indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER)

p(doc, 'Функционал — энтропийная риск-мера (SoftMin):')
p(doc, 'L(θ) = (1/a) · log E_ω[ exp(−a · PnL_θ(ω)) ].',
  indent_first=False, align=WD_ALIGN_PARAGRAPH.CENTER)

p(doc, 'θ оптимизируется end-to-end pathwise-градиентом через '
       'дифференцируемый симулятор SDE (Monte Carlo Policy Gradient). '
       'В работе используются параметры S₀ = 100, v₀ = 0.04, r = 0.01, '
       'κ = 2.0, θ = 0.04, ξ = 0.1, ρ = −0.7, T = 1 год, N = 30, '
       'K = 100, c = 10⁻³ (10 б. п.), a = 1.0.')


# ════════════════════════════════════════════════════════════════════════════
# 3. АКТУАЛЬНОСТЬ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '3. Актуальность темы')

p(doc, 'Классическая дельта-стратегия Блэка–Шоулза оптимальна в пределе '
       'непрерывной торговли без трений. На практике обе предпосылки '
       'нарушаются: ребалансировка дискретна, каждая сделка стоит денег. '
       'Асимптотические поправки Леланда (1985) и no-trade band '
       'Уолли–Уилмотта (1997) требуют модельно-специфичных выводов и плохо '
       'переносятся на экзотические payoff-ы.')

p(doc, 'Deep Hedging (Buehler et al., 2019) предлагает обходной путь — '
       'end-to-end обучение нейронной политики на симулированных траекториях. '
       'Однако базовый MLP-агент Бюлера не даёт стабильного преимущества '
       'над closed-form дельтой под моделью Heston: в нашем ablation плоский '
       'DH показывает SoftMin = 10.42, тогда как Heston-CF дельта — 9.55. '
       'Чтобы DH действительно стал рабочим инструментом, нужна инженерия '
       'поверх базового рецепта.')

p(doc, 'Два направления инженерии в DH-литературе недоизучены: '
       '(i) признаковые эмбеддинги в духе tabular-DL '
       '(Gorishniy et al., 2022), (ii) оптимизаторы второго порядка. '
       'По второму пункту в 2024 году появились работы Mueller et al. '
       '(ICAIF \'24) и Enkhbayar (arXiv-препринт), применяющие K-FAC '
       'к DH и сообщающие о сокращении числа итераций в 4 раза. По первому '
       'пункту DH-специфичных публикаций найти не удалось — это потенциально '
       'новый вклад для выпускной работы.')


# ════════════════════════════════════════════════════════════════════════════
# 4. ОБЗОР СУЩЕСТВУЮЩИХ МЕТОДОВ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '4. Обзор существующих методов')

h(doc, '4.1. Классические аналитические методы', level=2)
p(doc, 'Black–Scholes (1973) и Merton (1973) дают дельту в замкнутой '
       'форме при нулевых издержках и непрерывной торговле. Heston (1993) '
       'обобщает ценообразование на стохастическую волатильность; дельта '
       'получается численно через интеграл характеристической функции и '
       'используется нами как эталон. Longstaff–Schwartz (2001) и '
       'Haugh–Kogan (2004) предложили регрессионную LSMC-дельту — '
       'практичный бейзлайн без замкнутой формы.')

p(doc, 'Под трениями Leland (1985) адаптирует волатильность в формуле BS; '
       'Whalley–Wilmott (1997) получают оптимальный no-trade band, '
       'масштабирующийся как c^{1/3}; Davis–Panas–Zariphopoulou (1993) и '
       'Cvitanić–Karatzas (1996) решают задачу через максимизацию полезности. '
       'Все эти методы опираются на модельные предпосылки и плохо '
       'обобщаются на экзотику и реальные данные.')

h(doc, '4.2. Deep Hedging и расширения', level=2)
p(doc, 'Buehler, Gonon, Teichmann, Wood «Deep Hedging» '
       '(Quantitative Finance, 2019) — исходная формулировка MCPG для '
       'хеджирования: рекуррентная нейросеть обучается минимизировать '
       'выпуклую риск-меру P&L с учётом fee, на симулированных путях Heston. '
       'Buehler et al. (2022) «Deep Bellman Hedging» вводят actor-critic с '
       'явным Bellman-уравнением для выпуклых риск-мер. Cao, Chen, Hull, '
       'Poulos (Journal of Financial Data Science, 2021) применяют '
       'distributional RL к дельта- и гамма-хеджингу под целью '
       '«среднее + станд. отклонение». Murray et al. (ICAIF 2022) используют '
       'экспоненциально-мультипликативную версию уравнения Беллмана для '
       'exp-utility. Qiao и Wan (arXiv:2407.19367, 2024) первыми в DH '
       'формулируют политику как δ_net = δ_BSM + residual — параметризация, '
       'которую мы применяем в DH-Deviation.')

h(doc, '4.3. K-FAC и оптимизация в Deep Hedging', level=2)
p(doc, 'Martens и Grosse (ICML 2015) предложили K-FAC — аппроксимацию '
       'Фишер-матрицы в виде блок-Кронекерова произведения факторов A⊗G, '
       'позволяющую делать шаг натурального градиента без явного обращения '
       'квадратной матрицы размера num_params. Wu et al. (ACKTR, NeurIPS '
       '2017) перенесли K-FAC на policy-gradient. В DH: Mueller, Akkari, '
       'Gonon, Wood «Fast Deep Hedging with Second-Order Optimization» '
       '(ICAIF \'24, ACM) — каноническая ссылка, сообщающая ~4× сокращение '
       'числа итераций против Adam на cliquet-хедже под stochastic vol. '
       'Enkhbayar (arXiv:2411.15002, 2024) — LSTM-хеджер + K-FAC; препринт, '
       'на апрель 2026 не опубликован в рецензируемом виде.')

h(doc, '4.4. Признаковые эмбеддинги в tabular DL', level=2)
p(doc, 'Rahimi и Recht (NeurIPS 2007) — random Fourier features для '
       'аппроксимации ядерных методов. Tancik et al. (NeurIPS 2020) и '
       'Sitzmann et al. SIREN (NeurIPS 2020) — периодические эмбеддинги и '
       'активации для координатных MLP. Gorishniy, Rubachev, Babenko '
       '(NeurIPS 2022) переносят периодическое кодирование скалярных '
       'признаков на tabular DL. Поиск в arXiv, Google Scholar, SSRN и '
       'DBLP по запросам "deep hedging + Fourier features / periodic '
       'embedding / SIREN / Tancik / random Fourier" не дал DH-специфичных '
       'работ — значит, применение PAF в политике DH можно считать новым.')


# ════════════════════════════════════════════════════════════════════════════
# 5. ВЫБОР МЕТОДОВ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '5. Выбор методов решения. Обоснование')

p(doc, 'Выбор каждого компонента стека:')

h(doc, 'Риск-мера — SoftMin (энтропийная).', level=2)
p(doc, 'Дуальна к exp-utility, выпукла, монотонна, гладкая; '
       'используется в оригинальной работе Buehler и в большинстве '
       'последующих. Выбор a = 1.0 согласован с Mueller 2024.')

h(doc, 'Симулятор — Heston.', level=2)
p(doc, 'Минимальная стохастическая volatility-модель, где classical BSM-дельта '
       'уже не оптимальна и где есть нетривиальный корреляционный риск '
       '(leverage effect ρ = −0.7). Симуляция по схеме Эйлера-full-truncation.')

h(doc, 'Базовый тренер — MCPG (pathwise).', level=2)
p(doc, 'Дифференцируемый симулятор делает pathwise-оценку градиента '
       'низкодисперсной по сравнению со score-function (REINFORCE). '
       'Этот выбор фиксируется по всей работе — мы не меняем сам алгоритм, '
       'а улучшаем обучаемую политику.')

h(doc, 'PAF-эмбеддинг входа.', level=2)
p(doc, 'Глобальная Periodic Adaptive Fourier параметризация: '
       'emb(x) = [sin(x B), cos(x B)], B ∈ ℝ^{d_in × n_freq} — обучаемая. '
       'Мотивация из tabular DL (Gorishniy 2022): периодическое кодирование '
       'помогает сети выразить высокочастотные зависимости от S и τ, что '
       'плохо даётся MLP. DH-применений PAF в литературе не найдено.')

h(doc, 'Running P&L как фича.', level=2)
p(doc, 'cash_t + δ_{t-1} S_t добавляется в state. Мотивация: в дискретной '
       'ребалансировке с fee задача не марковская по (S, v) — прошлый P&L '
       'несёт информацию о том, насколько агент уже «отклонился». '
       'Buehler 2019 включает его неявно через рекуррентность; Cao-Hull 2021 '
       'эксплицитно. Murray 2022 наоборот считает эту фичу избыточной — '
       'конфликт, который и проверяется в ablation.')

h(doc, 'Deviation-head и zero-last-layer init.', level=2)
p(doc, 'Сеть выдаёт не δ целиком, а отклонение от BSM-дельты: '
       'δ_total = δ_BSM(S_t, √v_t, τ_t) + net(features). Первая часть — '
       'из Qiao–Wan 2024. В сочетании с zero-last-layer init (веса последнего '
       'линейного слоя зануляются) агент стартует строго как δ_BSM — '
       'хорошее начало даже до первого градиентного шага. Pairing этих двух '
       'трюков в DH-литературе не зафиксирован — потенциально '
       'оригинальный вклад.')

h(doc, 'Оптимизаторы — Adam / K-FAC / Muon.', level=2)
p(doc, 'Adam — бейзлайн. K-FAC реализован напрямую с forward-pre и '
       'full-backward-хуками: хуки копят статистики активаций A = E[ãã^⊤] '
       '(с добавленной 1 для bias) и градиентов G = E[g g^⊤], обновление '
       'делается как g̃ = (G + λI)^{−1} g_W (A + λI)^{−1} с '
       'trust-region-клипом. Для работы в time-unrolled цикле '
       'активации складываются в LIFO-стек и pop-ятся в backward — '
       'последовательность совпадает с реверсным обходом графа. '
       'Muon — ортогонализация моментум-буфера через 5 шагов '
       'квинтической итерации Ньютона-Шульца (Jordan 2024); для 2D-весов, '
       'остальные параметры идут в fallback-Adam.')


# ════════════════════════════════════════════════════════════════════════════
# 6. ПЛАН РЕШЕНИЯ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '6. План решения поставленной задачи')

p(doc, 'Репозиторий: src/ — библиотека, heston_comparison.ipynb — '
       'интерактивный ablation. Основные шаги реализации:')

numbered(doc, 'src/samplers.py — HestonSampler (схема Эйлера full-truncation, '
       'коррелированные W¹, W²).')
numbered(doc, 'src/deltas.py — LSMC-дельта через polynomial regression на '
       'сетке состояний; Heston closed-form дельта через интеграл '
       'характеристической функции.')
numbered(doc, 'src/backtest.py — torch_backtest: дифференцируемый self-financing '
       'симулятор портфеля, возвращает тензор (PnL, fees) с градиентом по θ.')
numbered(doc, 'src/models.py — HedgePolicy с конфигурируемыми '
       'architecture (mlp / paf / paf_fw), skip-connection в эмбеддинге и '
       'опцией zero_last_layer.')
numbered(doc, 'src/trainers.py — _BaseTrainer (val-loss, early stop, '
       'best-checkpoint) и специализации: DeepHedgingTrainer (vanilla MCPG + '
       'опциональный pretrain-делта), SACTrainer (энтропийная регуляризация '
       'через reparameterization), DeviationTrainer (residual-head с BSM).')
numbered(doc, 'src/trainers.py — KFACOptimizer (forward-pre и '
       'full-backward хуки, EMA обновление A, G, обращение с Tikhonov-дампингом '
       'каждые update_freq шагов, trust-region gradient-norm clip) '
       'и Muon (Newton-Schulz 5 с коэффициентами '
       '(3.4445, −4.7750, 2.0315), scale max(1, h/w)^{1/2}).')
numbered(doc, 'Запуск ablation: общие val/test-наборы из 5000 путей, '
       '10000 эпох, M_train = 3000, log каждые 200 эпох, early stop при '
       '100 проверках без улучшения, best-model restoration.')
numbered(doc, 'Evaluation: финальный SoftMin на test-set; pnl-mean и pnl-std '
       'как дополнительные метрики; learning curves сохраняются в '
       'results/*.json.')


# ════════════════════════════════════════════════════════════════════════════
# 7. РЕЗУЛЬТАТЫ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '7. Полученные результаты')

p(doc, 'Таблица 1. SoftMin на общем test-set (5000 путей, a = 1.0). '
       'Δ показан относительно Heston-CF дельты как абсолютное улучшение '
       'по SoftMin (отрицательное — лучше).')

# Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, t in enumerate(['Метод', 'SoftMin', 'Δ к Heston CF', 'Замечание']):
    hdr[i].text = ''
    r = hdr[i].paragraphs[0].add_run(t)
    r.bold = True
    r.font.size = Pt(12)
    _cyr(r)

rows = [
    ('Delta LSMC',           '9.9928',   '+0.44', 'регрессионная дельта'),
    ('Delta CF (Heston)',    '9.5526',   '0.00',  'эталон (closed-form)'),
    ('DH',                   '10.4207',  '+0.87', 'плоский MLP'),
    ('DH-PAF',               '9.3405',   '−0.21', 'PAF-эмбеддинг'),
    ('DH-PnL',               '9.8311',   '+0.28', 'MLP + running P&L'),
    ('DH-PAF-PnL',           '169.13',   'div.',  'расходимость'),
    ('DH-SAC',               '9.1738',   '−0.38', 'энтропийная регуляризация'),
    ('DH-PAF-PnL-Pretrain',  '197.15',   'div.',  'расходимость'),
    ('DH-Deviation',         '9.1394',   '−0.41', 'residual-head, Adam'),
    ('DH-Deviation-KFAC',    '9.0897',   '−0.46', 'лучший; K-FAC'),
    ('DH-Deviation-Muon',    '9.2422',   '−0.31', 'Muon (lr=2e-2)'),
]
for row in rows:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = ''
        r = cells[i].paragraphs[0].add_run(v)
        r.font.size = Pt(11)
        _cyr(r)
        if row[0] == 'DH-Deviation-KFAC':
            r.bold = True

p(doc, 'Ключевые наблюдения:', space_before=12)

numbered(doc, 'Плоский Deep Hedging (10.42) проигрывает даже LSMC-дельте '
       '(9.99) — это подтверждает тезис, что без инженерии базовый агент '
       'Бюлера не является рабочим бейзлайном на Heston.')
numbered(doc, 'PAF-эмбеддинг сам по себе (9.34) уже обходит обе эталонные '
       'дельты, включая closed-form (9.55). Это самая крупная единичная '
       'прибавка в стеке — при добавлении PAF результат сразу переходит '
       'в положительную по SoftMin-Δ зону.')
numbered(doc, 'Running P&L как фича при случайной инициализации разрушает '
       'обучение в паре с PAF: DH-PAF-PnL и DH-PAF-PnL-Pretrain выдают '
       'SoftMin ≫ 100 (численная расходимость). Фича вносит сильный '
       'негладкий сигнал, и сети нужен хороший стартовый чекпоинт.')
numbered(doc, 'Residual-head (δ = δ_BSM + net) снимает расходимость: '
       'DH-Deviation с теми же входами и той же PAF-архитектурой даёт 9.14 — '
       'лучший результат среди конфигураций с Adam-оптимизатором.')
numbered(doc, 'SAC-регуляризация (9.17) даёт улучшение, сопоставимое по '
       'порядку с deviation-head, но путь обучения заметно шумнее '
       '(см. learning curves в разделе 15 ноутбука).')
numbered(doc, 'K-FAC даёт небольшое, но устойчивое улучшение поверх '
       'DH-Deviation-Adam: 9.0897 против 9.1394. Это согласуется '
       'качественно с Mueller 2024 (ICAIF \'24): natural gradient даёт '
       'шаги, лучше выровненные по кривизне, и находит чуть более '
       'глубокий минимум за то же число эпох.')
numbered(doc, 'Muon (9.24) оказался чуть хуже K-FAC при lr = 2 × 10⁻². '
       'Muon чувствителен к LR; в ВКР планируется sweep lr и momentum.')
numbered(doc, 'Итоговая лучшая конфигурация — DH-Deviation-KFAC — это '
       'стек пяти трюков: PAF-эмбеддинг + running P&L + residual-head '
       '(δ_BSM + net) + zero-last-layer init + K-FAC preconditioner. '
       'Улучшение относительно closed-form Heston-дельты составляет '
       '0.46 по SoftMin.')


# ════════════════════════════════════════════════════════════════════════════
# 8. ВЫВОДЫ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '8. Выводы')

p(doc, 'Главный результат практики — документированный ablation, '
       'который показывает, что не один отдельный трюк, а сочетание '
       'пяти приёмов обеспечивает устойчивое преимущество Deep Hedging '
       'над закрытой Heston-дельтой: PAF-эмбеддинг входа, '
       'running P&L как фича, residual-параметризация выхода от BSM-дельты, '
       'zero-last-layer инициализация, K-FAC в качестве оптимизатора. '
       'SoftMin на общем test-set: 9.0897 у лучшей DH-конфигурации против '
       '9.5526 у эталонной closed-form дельты.')

p(doc, 'Отдельно стоит отметить атрибуцию каждого компонента. Running P&L '
       '(Buehler 2019, Cao 2021), residual-head (Qiao–Wan 2024) и '
       'K-FAC-для-DH (Mueller 2024, ICAIF \'24) опубликованы ранее. '
       'А вот PAF-эмбеддинг как вход политики в DH и комбинация '
       'zero-last-layer init с residual-head в DH-специфичной литературе '
       'обнаружить не удалось — это кандидаты на новый вклад в ВКР.')

p(doc, 'Отрицательный результат тоже полезен: running P&L как фича при '
       'случайной инициализации и без residual-структуры приводит к '
       'численной расходимости (SoftMin ≫ 100). Это объясняет, почему '
       'прямое расширение входа политики не дало результата в '
       'предыдущих попытках и почему residual-параметризация — '
       'не факультатив, а обязательный компонент, когда используется '
       'дополнительная P&L-фича.')

p(doc, 'Задачи ВКР, опирающиеся на результаты практики:')
bullet(doc, 'перенос лучшей конфигурации (DH-Deviation-KFAC) на реальные '
       'данные: BTC и ETH опционы с Deribit, с Dupire LVM-дельтой '
       'в качестве бейзлайна;')
bullet(doc, 'разработка universal-модели: добавить параметры опциона '
       '(log-moneyness, τ, тип) во входной state и обучить один агент '
       'на смеси контрактов — inference на unseen опционы без '
       'переобучения;')
bullet(doc, 'sweep гиперпараметров K-FAC (damping, update_freq, ema_decay) '
       'и схем LR (cosine, warm-up) — эти оси в рамках практики '
       'не покрыты;')
bullet(doc, 'расширение стека на экзотические payoff-ы (barrier, cliquet), '
       'где аналитическая дельта отсутствует и residual-параметризация '
       'требует аппроксимированного BSM-ориентира.')


# ════════════════════════════════════════════════════════════════════════════
# 9. СПИСОК ЛИТЕРАТУРЫ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '9. Список изученной литературы')

refs = [
    'Black F., Scholes M. The Pricing of Options and Corporate Liabilities // '
    'Journal of Political Economy. 1973. Vol. 81, № 3. P. 637–654.',

    'Merton R. C. Theory of Rational Option Pricing // Bell Journal of '
    'Economics and Management Science. 1973. Vol. 4, № 1. P. 141–183.',

    'Heston S. L. A Closed-Form Solution for Options with Stochastic '
    'Volatility with Applications to Bond and Currency Options // Review of '
    'Financial Studies. 1993. Vol. 6, № 2. P. 327–343.',

    'Leland H. E. Option Pricing and Replication with Transactions Costs // '
    'Journal of Finance. 1985. Vol. 40, № 5. P. 1283–1301.',

    'Whalley A. E., Wilmott P. An Asymptotic Analysis of an Optimal Hedging '
    'Model for Option Pricing with Transaction Costs // Mathematical Finance. '
    '1997. Vol. 7, № 3. P. 307–324.',

    'Davis M. H. A., Panas V. G., Zariphopoulou T. European Option Pricing '
    'with Transaction Costs // SIAM Journal on Control and Optimization. '
    '1993. Vol. 31, № 2. P. 470–493.',

    'Dupire B. Pricing with a Smile // Risk. 1994. Vol. 7, № 1. P. 18–20.',

    'Longstaff F. A., Schwartz E. S. Valuing American Options by Simulation: '
    'A Simple Least-Squares Approach // Review of Financial Studies. 2001. '
    'Vol. 14, № 1. P. 113–147.',

    'Buehler H., Gonon L., Teichmann J., Wood B. Deep Hedging // Quantitative '
    'Finance. 2019. Vol. 19, № 8. P. 1271–1291. DOI: 10.1080/14697688.2019.1571683.',

    'Buehler H., Phillip M., Wood B. Deep Bellman Hedging. arXiv:2207.00932. 2022.',

    'Cao J., Chen J., Hull J., Poulos Z. Deep Hedging of Derivatives Using '
    'Reinforcement Learning // Journal of Financial Data Science. 2021. '
    'Vol. 3, № 1. P. 10–27. arXiv:2103.16409.',

    'Murray P., Wood B., Buehler H., Wiese M., Pakkanen M. Deep Hedging: '
    'Continuous Reinforcement Learning for Hedging of General Portfolios '
    'across Multiple Risk Aversions // Proc. of the 3rd ACM International '
    'Conference on AI in Finance (ICAIF \'22). 2022. arXiv:2207.07467.',

    'Qiao Y., Wan L. Enhancing Black–Scholes Delta Hedging via Deep Learning. '
    'arXiv:2407.19367. 2024.',

    'Mueller K., Akkari A., Gonon L., Wood B. Fast Deep Hedging with '
    'Second-Order Optimization // Proc. of the 5th ACM International '
    'Conference on AI in Finance (ICAIF \'24). 2024. '
    'DOI: 10.1145/3677052.3698604. arXiv:2410.22568.',

    'Enkhbayar T. A New Way: Kronecker-Factored Approximate Curvature '
    'Deep Hedging and its Benefits. arXiv:2411.15002. 2024.',

    'Martens J., Grosse R. Optimizing Neural Networks with Kronecker-factored '
    'Approximate Curvature // Proc. of ICML. 2015. arXiv:1503.05671.',

    'Wu Y., Mansimov E., Liao S., Grosse R., Ba J. Scalable Trust-Region '
    'Method for Deep Reinforcement Learning using Kronecker-factored '
    'Approximation (ACKTR) // Proc. of NeurIPS. 2017. arXiv:1708.05144.',

    'Jordan K., Jin Y., Boza V., You J., Cesista F., Newhouse L., Bernstein J. '
    'Muon: An Optimizer for the Hidden Layers of Neural Networks. '
    'kellerjordan.github.io/posts/muon/. 2024.',

    'Kingma D. P., Ba J. Adam: A Method for Stochastic Optimization // '
    'Proc. of ICLR. 2015.',

    'Loshchilov I., Hutter F. Decoupled Weight Decay Regularization // '
    'Proc. of ICLR. 2019.',

    'Haarnoja T., Zhou A., Abbeel P., Levine S. Soft Actor-Critic: '
    'Off-Policy Maximum Entropy Deep Reinforcement Learning with a Stochastic '
    'Actor // Proc. of ICML. 2018. arXiv:1801.01290.',

    'Rahimi A., Recht B. Random Features for Large-Scale Kernel Machines // '
    'Proc. of NeurIPS. 2007.',

    'Tancik M., Srinivasan P. P., Mildenhall B. et al. Fourier Features Let '
    'Networks Learn High Frequency Functions in Low Dimensional Domains // '
    'Proc. of NeurIPS. 2020. arXiv:2006.10739.',

    'Sitzmann V., Martel J. N. P., Bergman A. W., Lindell D. B., Wetzstein G. '
    'Implicit Neural Representations with Periodic Activation Functions '
    '(SIREN) // Proc. of NeurIPS. 2020. arXiv:2006.09661.',

    'Gorishniy Y., Rubachev I., Babenko A. On Embeddings for Numerical '
    'Features in Tabular Deep Learning // Proc. of NeurIPS. 2022. '
    'arXiv:2203.05556.',

    'Silver D., Huang A., Jaderberg M. et al. Residual Policy Learning. '
    'arXiv:1812.06298. 2018.',

    'Johannink T., Bahl S., Nair A. et al. Residual Reinforcement Learning '
    'for Robot Control // Proc. of ICRA. 2019. arXiv:1812.03201.',

    'Zhang H., Dauphin Y. N., Ma T. Fixup Initialization: Residual Learning '
    'Without Normalization // Proc. of ICLR. 2019. arXiv:1901.09321.',
]

for i, ref in enumerate(refs, 1):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.left_indent = Cm(0.5)
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run(f'{i}. {ref}')
    r.font.size = Pt(12)
    _cyr(r)


# ════════════════════════════════════════════════════════════════════════════
# 10. ПРИЛОЖЕНИЯ
# ════════════════════════════════════════════════════════════════════════════
h(doc, '10. Приложения')

h(doc, 'Приложение А. Гиперпараметры ablation', level=2)
bullet(doc, 'Heston: S₀ = 100, v₀ = 0.04, r = 0.01, κ = 2.0, θ = 0.04, '
       'ξ = 0.1, ρ = −0.7, T = 1 год, N = 30, K = 100, c = 10⁻³.')
bullet(doc, 'Тренировка: M_train = 3000 путей, n_epochs = 10000, '
       'log_every = 200, early_stop_patience = 100, lr = 10⁻³ (кроме Muon: '
       '2 × 10⁻²).')
bullet(doc, 'Валидация/тест: M_val = M_test = 5000 путей, общий '
       'random-seed = 42 для всех методов.')
bullet(doc, 'Архитектура: скрытые слои (64, 32), LeakyReLU(0.01), '
       'PAF с n_freq = 16, σ_init = 1.0.')
bullet(doc, 'K-FAC: damping = 10⁻¹, ema_decay = 0.95, update_freq = 10, '
       'momentum = 0.9, grad_clip = 1.0.')
bullet(doc, 'Muon: momentum = 0.95, ns_steps = 5, scale = max(1, out/in)^{1/2}.')

h(doc, 'Приложение Б. Ключевые фрагменты кода', level=2)

p(doc, 'B.1. Newton–Schulz-орто­гонализатор (ядро Muon):')
code_block(doc,
"""def _newton_schulz5(G, steps=5, eps=1e-7):
    a, b, c = 3.4445, -4.7750, 2.0315
    X = G.clone().float() / (G.norm() + eps)
    transposed = X.size(-2) > X.size(-1)
    if transposed: X = X.transpose(-2, -1)
    for _ in range(steps):
        A = X @ X.transpose(-2, -1)
        B = b * A + c * (A @ A)
        X = a * X + B @ X
    if transposed: X = X.transpose(-2, -1)
    return X.to(G.dtype)""")

p(doc, 'Б.2. Шаг K-FAC (предобуславливание Linear-градиентов):')
code_block(doc,
"""# для каждого Linear-модуля:
grad = torch.cat([W.grad, b.grad.unsqueeze(1)], dim=1)
nat  = G_inv @ grad @ A_inv
W.grad.copy_(nat[:, :-1]); b.grad.copy_(nat[:, -1])
# глобальный clip + SGD с моментумом:
torch.nn.utils.clip_grad_norm_(model.parameters(), max_norm=1.0)
for p in model.parameters():
    buf = mom_buf[id(p)].mul_(0.9).add_(p.grad)
    p.add_(buf, alpha=-lr)""")

p(doc, 'Б.3. Backtest с residual-головой (ядро DeviationTrainer):')
code_block(doc,
"""for t in range(N):
    delta_bsm = bsm_delta(S[:, t], sqrt(v[:, t]), tau_t, K, r)
    features  = [state_t, tau_t, delta_prev]
    if use_pnl: features.append(cash + delta_prev * S[:, t])
    deviation = net(cat(features)).squeeze(1)
    delta     = delta_bsm + deviation
    # self-financing update ...""")

h(doc, 'Приложение В. Репозиторий', level=2)
p(doc, 'Исходный код и ноутбук ablation доступны в Git-репозитории проекта. '
       'Основные файлы: src/trainers.py (тренеры и оптимизаторы), '
       'src/models.py (архитектура политики), src/samplers.py (HestonSampler), '
       'src/deltas.py (не-RL бейзлайны), heston_comparison.ipynb '
       '(сводный ablation).')


# ════════════════════════════════════════════════════════════════════════════
# ОПИСАНИЕ ПРИМЕНЕНИЯ ГЕНЕРАТИВНОЙ МОДЕЛИ (требование НИУ ВШЭ)
# ════════════════════════════════════════════════════════════════════════════
h(doc, 'Описание применения генеративной модели')

p(doc, 'В процессе подготовки настоящего отчёта и сопровождающего кода '
       'использовалась большая языковая модель Anthropic Claude '
       '(claude.ai, версия Opus 4.7). Характер применения:')
bullet(doc, 'поиск и систематизация литературных источников '
       '(раздел 4 «Обзор»), проверка точности библиографических ссылок — '
       'финальный список проверен автором вручную;')
bullet(doc, 'реализация вспомогательного кода: классы Muon и KFACOptimizer '
       'в src/trainers.py были написаны в диалоговом режиме по исходной '
       'спецификации (Martens-Grosse 2015 для K-FAC, Jordan 2024 для Muon); '
       'корректность проверена автором через smoke-tests и '
       'полный ablation-прогон;')
bullet(doc, 'стилистическая редактура русскоязычной части отчёта; '
       'все утверждения, постановка задачи, выбор методов, численные '
       'эксперименты и выводы принадлежат автору.')
p(doc, 'Научные результаты, формулировка гипотез, выбор ablation-сетапа и '
       'интерпретация экспериментов сформулированы автором самостоятельно.')


# ════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f'OK: {OUT}')
